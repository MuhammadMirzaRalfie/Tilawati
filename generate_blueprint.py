"""
Blueprint Generator - TilawatiMobile
Generates blueprint_TilawatiMobile.docx based on the project state and progress.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ COLORS ============
PRIMARY = RGBColor(0x1B, 0x5E, 0x20)        # Dark green
SECONDARY = RGBColor(0x0D, 0x47, 0xA1)      # Dark blue
ACCENT = RGBColor(0xBF, 0x36, 0x0C)         # Orange
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "E8F5E9"

# ============ HELPERS ============

def set_cell_bg(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tc_pr.append(shd)

def heading(doc, text, level=1, color=PRIMARY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h

def para(doc, text, bold=False, italic=False, size=11, color=None, align=None, justify=True):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg="1B5E20", header_white=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        if header_white:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if header_bg:
            set_cell_bg(hdr_cells[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(c))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t

def page_break(doc):
    doc.add_page_break()

def hr(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    bdr.append(bottom)
    p_pr.append(bdr)


# ============ DOCUMENT ============
doc = Document()

# Default styles
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

# Margins
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Configure heading styles
for lvl in range(1, 4):
    s = doc.styles[f'Heading {lvl}']
    s.font.name = 'Calibri'
    s.font.color.rgb = PRIMARY
    if lvl == 1:
        s.font.size = Pt(20)
    elif lvl == 2:
        s.font.size = Pt(15)
    else:
        s.font.size = Pt(13)

# ============ COVER PAGE ============
for _ in range(4):
    doc.add_paragraph()

t1 = doc.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t1.add_run("BLUEPRINT PENGEMBANGAN")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = PRIMARY
r.font.name = 'Calibri'

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run("TilawatiMobile")
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = SECONDARY
r.font.name = 'Calibri'

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Aplikasi Pembelajaran Bacaan Al-Qur'an Berbasis AI\ndengan Metode Tilawati Jilid 1")
r.font.size = Pt(14)
r.italic = True
r.font.color.rgb = GRAY
r.font.name = 'Calibri'

for _ in range(3):
    doc.add_paragraph()

# Document info table
info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Light Grid Accent 1'
info_data = [
    ("Nama Proyek", "TilawatiMobile (Tugas Akhir)"),
    ("Judul TA", "Tilawati Hijaiyah ASR — Sistem Pengenalan Ucapan 28 Huruf Hijaiyah"),
    ("Versi Dokumen", "1.0"),
    ("Tanggal", "30 April 2026"),
    ("Status", "Active Development — Fase Aplikasi Mobile (Tahap Akhir)"),
    ("Komponen Utama", "Flutter (Mobile) + FastAPI (Backend) + PostgreSQL + Model HPT-D"),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = ''
    p1 = info_table.rows[i].cells[0].paragraphs[0]
    rk = p1.add_run(k)
    rk.bold = True
    rk.font.size = Pt(10)
    set_cell_bg(info_table.rows[i].cells[0], "F1F8E9")
    info_table.rows[i].cells[1].text = ''
    p2 = info_table.rows[i].cells[1].paragraphs[0]
    rv = p2.add_run(v)
    rv.font.size = Pt(10)

for _ in range(2):
    doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Dokumen Internal — Rancangan Pengembangan End-to-End")
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = GRAY

page_break(doc)

# ============ DAFTAR ISI ============
heading(doc, "Daftar Isi", level=1)
toc_entries = [
    ("1. Ringkasan Eksekutif", ""),
    ("2. Latar Belakang & Konteks", ""),
    ("3. Tujuan, Sasaran, dan Persona Pengguna", ""),
    ("4. Ruang Lingkup", ""),
    ("5. Arsitektur Sistem End-to-End", ""),
    ("6. Komponen AI/ML — Model HPT-D dan Klasifikasi", ""),
    ("7. Arsitektur Backend (FastAPI)", ""),
    ("8. Arsitektur Frontend (Flutter)", ""),
    ("9. Spesifikasi Fungsional", ""),
    ("10. Spesifikasi Non-Fungsional", ""),
    ("11. Data Model & Database Schema", ""),
    ("12. API Specification", ""),
    ("13. User Flow & Skenario Penggunaan", ""),
    ("14. Roadmap Pengembangan", ""),
    ("15. Status Progres Saat Ini (30 April 2026)", ""),
    ("16. Strategi Testing", ""),
    ("17. Strategi Deployment", ""),
    ("18. Risiko & Mitigasi", ""),
    ("19. Pengembangan Masa Depan", ""),
    ("20. Lampiran", ""),
]
for entry, _ in toc_entries:
    p = doc.add_paragraph(entry)
    p.paragraph_format.left_indent = Inches(0.25)
    for run in p.runs:
        run.font.size = Pt(11)

page_break(doc)

# ============ 1. RINGKASAN EKSEKUTIF ============
heading(doc, "1. Ringkasan Eksekutif", level=1)

heading(doc, "1.1 Visi", level=2)
para(doc,
    "TilawatiMobile adalah aplikasi pembelajaran Al-Qur'an berbasis AI yang dirancang untuk membantu "
    "pemula memperbaiki bacaan huruf hijaiyah berharakat fathah secara mandiri, mengikuti kurikulum "
    "Tilawati Jilid 1. Aplikasi ini menggabungkan model Automatic Speech Recognition (ASR) khusus "
    "huruf hijaiyah dengan antarmuka mobile yang ramah anak/pemula sehingga proses koreksi bacaan "
    "yang biasanya bergantung pada kehadiran ustadz dapat berlangsung kapan saja, di mana saja."
)

heading(doc, "1.2 Permasalahan", level=2)
bullet(doc, "Pembelajaran Tilawati konvensional bergantung penuh pada kehadiran ustadz/ustadzah untuk koreksi langsung.")
bullet(doc, "Anak/pemula yang belajar mandiri tidak mendapatkan feedback akurat terkait makharijul huruf (titik artikulasi).")
bullet(doc, "Buku Tilawati hanya menyajikan teks Arab + transliterasi, tanpa mekanisme verifikasi pelafalan.")
bullet(doc, "Aplikasi belajar Al-Qur'an di pasaran umumnya menggunakan model ASR umum berbahasa Arab — tidak fokus pada 28 huruf hijaiyah dengan harakat fathah saja, sehingga akurasi rendah untuk pemula.")

heading(doc, "1.3 Solusi", level=2)
para(doc,
    "TilawatiMobile mengadopsi pendekatan tiga lapis: (1) model ASR HPT-D yang dilatih khusus untuk "
    "rangkaian 2-4 huruf hijaiyah dengan harakat fathah; (2) backend FastAPI yang menggabungkan "
    "evaluasi audio, manajemen progres, dan autentikasi pengguna; (3) aplikasi Flutter dengan UX "
    "rekam-per-kata sehingga setiap kelompok huruf dievaluasi langsung dengan visual feedback "
    "(hijau untuk benar, merah untuk salah)."
)

heading(doc, "1.4 Fitur Utama", level=2)
make_table(doc,
    ["Fitur", "Status", "Deskripsi Singkat"],
    [
        ("Autentikasi (Register/Login)", "Selesai", "JWT-based, password hashing dengan bcrypt"),
        ("Konten Jilid 1 (44 halaman)", "Selesai", "Teks Arab + transliterasi sesuai kurikulum Tilawati"),
        ("Konten Jilid 2-6", "Minimal", "Data dummy 2-4 pelajaran per jilid (perlu enrichment)"),
        ("Latihan Per-Kata", "Selesai", "Rekam satu grup huruf, evaluasi instant via HPT-D"),
        ("Visual Highlight Arabic", "Selesai", "Grup aktif disorot hijau dengan underline"),
        ("Glosarium 28 Huruf", "Selesai", "Klasifikasi per huruf via model anas"),
        ("Riwayat Evaluasi", "Selesai", "Pull-to-refresh, batch 20 item"),
        ("Dashboard Progres", "Selesai", "Stat card + progress bar per jilid"),
        ("Audio Referensi per Huruf", "Belum", "Direncanakan Fase 5"),
        ("Mode Offline Terbatas", "Belum", "Direncanakan Fase 5"),
        ("Deploy HuggingFace Spaces", "Belum", "Direncanakan Fase 4"),
    ],
    col_widths=[1.7, 0.9, 3.8],
)

heading(doc, "1.5 Status Saat Ini", level=2)
para(doc,
    "Per 30 April 2026, fase Klasifikasi (Phase 1) dan ASR (Phase 2) telah selesai dengan model final "
    "HPT-D mencapai CER 0.1013, WER 0.2025, dan TSA per-class accuracy 65.2%. Fase Mobile App + Backend "
    "(Phase 3) berada di tahap akhir: integrasi end-to-end Flutter → FastAPI → HPT-D telah berfungsi, "
    "44 halaman Jilid 1 lengkap, latihan per-kata dengan instant feedback aktif, history dan "
    "progress dashboard terhubung ke API real. Sisa pekerjaan adalah pengayaan konten Jilid 2-6, "
    "testing E2E menyeluruh, dan deployment ke HuggingFace Spaces."
)

page_break(doc)

# ============ 2. LATAR BELAKANG & KONTEKS ============
heading(doc, "2. Latar Belakang & Konteks", level=1)

heading(doc, "2.1 Konteks Akademik", level=2)
para(doc,
    "Aplikasi TilawatiMobile merupakan deliverable utama dari Tugas Akhir (TA) berjudul "
    "“Tilawati Hijaiyah ASR — Sistem Pengenalan Ucapan 28 Huruf Hijaiyah Berharakat Fathah”. "
    "TA ini terbagi dalam tiga fase besar: Klasifikasi (Phase 1), ASR (Phase 2), dan Deploy (Phase 3). "
    "Phase 1 dan 2 berfokus pada R&D model machine learning, sementara Phase 3 — yang menjadi fokus "
    "dokumen blueprint ini — bertujuan mengintegrasikan model ke produk yang dapat dipakai end-user."
)

heading(doc, "2.2 Metode Tilawati", level=2)
para(doc,
    "Tilawati adalah metode pembelajaran membaca Al-Qur'an yang dikembangkan di Surabaya, Indonesia, "
    "dan dipakai secara luas di TPQ/TPA seantero negeri. Metode ini menyusun materi dalam enam jilid "
    "berurutan:"
)
bullet(doc, "Jilid 1 — Pengenalan 28 huruf hijaiyah dengan harakat fathah (target aplikasi ini)")
bullet(doc, "Jilid 2 — Huruf sambung dan kata sederhana")
bullet(doc, "Jilid 3 — Hukum Mad (panjang) dan sukun")
bullet(doc, "Jilid 4 — Hukum Nun & Mim sukun (Idzhar, Idgham, Ikhfa, Iqlab)")
bullet(doc, "Jilid 5 — Waqaf dan Ibtida")
bullet(doc, "Jilid 6 — Bacaan gharib & musykilat")
para(doc,
    "Dalam metode Tilawati, ustadz mencontohkan bacaan terlebih dahulu (talqin), murid menirukan, "
    "dan ustadz mengoreksi makharijul huruf serta tajwid. TilawatiMobile menggantikan tahap koreksi "
    "ini dengan ASR — tetapi tidak menggantikan kebutuhan ustadz untuk pengayaan tajwid lanjutan."
)

heading(doc, "2.3 Tantangan Pembelajaran Konvensional", level=2)
make_table(doc,
    ["Tantangan", "Dampak"],
    [
        ("Ketergantungan pada ustadz", "Murid tidak bisa latihan saat ustadz tidak hadir"),
        ("Rasio guru-murid 1:20+ di TPQ", "Koreksi per anak terbatas waktu, sebagian huruf terlewat"),
        ("Tidak ada feedback objektif", "Anak ragu apakah bacaannya benar saat belajar di rumah"),
        ("Kualitas pengajaran bervariasi", "Murid pindah TPQ harus mengulang/menyesuaikan ulang"),
        ("Buku statis, tidak interaktif", "Engagement rendah untuk anak generasi digital"),
    ],
    col_widths=[2.0, 4.4],
)

heading(doc, "2.4 Peran Teknologi ASR", level=2)
para(doc,
    "Automatic Speech Recognition (ASR) modern berbasis model self-supervised (Wav2Vec2) mampu "
    "mempelajari representasi audio dengan dataset yang relatif kecil. Model HPT-D yang dipakai "
    "TilawatiMobile dilatih dari basis anas/wav2vec2-large-xlsr-arabic dan kemudian di-fine-tune "
    "dengan dataset rekaman 28 huruf hijaiyah berharakat fathah dalam bentuk rangkaian 2-4 huruf. "
    "Hasilnya adalah model spesialis hijaiyah yang lebih akurat untuk kasus penggunaan TPQ/TPA "
    "dibandingkan ASR Arab umum."
)

page_break(doc)

# ============ 3. TUJUAN ============
heading(doc, "3. Tujuan, Sasaran, dan Persona Pengguna", level=1)

heading(doc, "3.1 Tujuan Umum", level=2)
para(doc,
    "Membangun aplikasi mobile Android/iOS yang menggabungkan model ASR HPT-D dengan kurikulum "
    "Tilawati Jilid 1 sehingga pengguna dapat berlatih dan memperoleh feedback akurat secara mandiri "
    "tanpa mengorbankan validitas evaluasi pelafalan."
)

heading(doc, "3.2 Tujuan Khusus", level=2)
numbered(doc, "Menyediakan 44 halaman konten Jilid 1 dengan teks Arab dan transliterasi yang akurat sesuai buku Tilawati.")
numbered(doc, "Mengimplementasikan latihan per-kata (per grup huruf) dengan feedback instant: rekam → evaluasi → next.")
numbered(doc, "Menyajikan dashboard progres yang menampilkan riwayat latihan, skor rata-rata, dan grade.")
numbered(doc, "Menyediakan glosarium 28 huruf dengan klasifikasi audio per huruf untuk reinforcement.")
numbered(doc, "Menjamin akurasi evaluasi ≥ 65% untuk per-class TSA berdasarkan model HPT-D.")
numbered(doc, "Menjamin response time evaluasi per kata ≤ 3 detik pada koneksi LAN.")

heading(doc, "3.3 Persona Pengguna", level=2)

heading(doc, "Persona 1: Anak TPQ Mandiri (Aisyah, 8 tahun)", level=3)
bullet(doc, "Sudah ikut kelas TPQ sore hari, ingin latihan di rumah dipantau orangtua.")
bullet(doc, "Tidak nyaman direkam suara nyaring di kelas. Lebih percaya diri latihan di rumah.")
bullet(doc, "Kebutuhan: aplikasi yang mudah dipakai, indikator visual jelas, animasi/warna menarik.")

heading(doc, "Persona 2: Pemula Dewasa (Pak Budi, 35 tahun)", level=3)
bullet(doc, "Mualaf, baru belajar Al-Qur'an. Malu belajar di kelas anak-anak.")
bullet(doc, "Punya smartphone Android, sering belajar lewat YouTube tapi tanpa feedback.")
bullet(doc, "Kebutuhan: aplikasi yang men-track progres, feedback objektif, bisa diulang sebanyak yang diperlukan.")

heading(doc, "Persona 3: Orangtua Pengawas (Ibu Sinta, 32 tahun)", level=3)
bullet(doc, "Anaknya kelas 2 SD, sedang belajar Tilawati Jilid 1.")
bullet(doc, "Ingin tahu progres anak tanpa harus tatap muka dengan ustadz tiap minggu.")
bullet(doc, "Kebutuhan: dashboard progres anak, riwayat latihan, indikator jilid yang sudah dikuasai.")

page_break(doc)

# ============ 4. RUANG LINGKUP ============
heading(doc, "4. Ruang Lingkup", level=1)

heading(doc, "4.1 In-Scope (Tercakup)", level=2)
bullet(doc, "Aplikasi mobile Flutter untuk Android (prioritas), iOS, dan Web (sekunder).")
bullet(doc, "Backend FastAPI dengan database PostgreSQL.")
bullet(doc, "Model ASR HPT-D untuk evaluasi pelafalan 28 huruf hijaiyah berharakat fathah (rangkaian 2-4 huruf).")
bullet(doc, "Model Klasifikasi anas untuk glosarium per-huruf (1 huruf per audio).")
bullet(doc, "Konten Tilawati Jilid 1 lengkap (44 halaman).")
bullet(doc, "Konten Jilid 2-6 dalam bentuk minimal/teaser (2-4 pelajaran per jilid).")
bullet(doc, "Autentikasi user via JWT.")
bullet(doc, "Manajemen progres per user per jilid.")
bullet(doc, "Riwayat evaluasi dengan filter per jilid.")

heading(doc, "4.2 Out-of-Scope (Tidak Tercakup)", level=2)
bullet(doc, "Konten Tilawati Jilid 2-6 lengkap (akan dikembangkan setelah TA).")
bullet(doc, "Hukum tajwid lanjutan (idzhar, idgham, ikhfa, iqlab) — model HPT-D belum dilatih untuk konteks ini.")
bullet(doc, "Pengenalan harakat selain fathah (kasrah, dhammah, sukun, tasydid, mad).")
bullet(doc, "Sistem ujian / sertifikasi formal.")
bullet(doc, "Integrasi dengan sistem TPQ/lembaga formal.")
bullet(doc, "Live class / video conference dengan ustadz.")
bullet(doc, "Mode offline penuh (offline terbatas baru direncanakan di Fase 5).")
bullet(doc, "Push notifikasi (direncanakan di Fase 5).")

heading(doc, "4.3 Asumsi", level=2)
bullet(doc, "Pengguna memiliki smartphone Android 7.0+ atau iOS 12+ dengan mikrofon berfungsi.")
bullet(doc, "Pengguna memiliki koneksi internet minimal 1 Mbps untuk evaluasi audio.")
bullet(doc, "Lingkungan rekam relatif tenang (noise floor < 50dB).")
bullet(doc, "Pengguna mengucapkan huruf dengan harakat fathah saja (sesuai Jilid 1).")
bullet(doc, "Backend dapat diakses melalui jaringan lokal (Docker compose) atau HuggingFace Spaces.")

heading(doc, "4.4 Batasan Sistem", level=2)
make_table(doc,
    ["Aspek", "Batasan"],
    [
        ("Audio input", "16 kHz mono WAV, durasi 0.3-10 detik per rekaman"),
        ("Cakupan huruf", "28 huruf hijaiyah dengan harakat fathah saja"),
        ("Panjang rekaman", "1-4 huruf per rekaman (sesuai data training HPT-D)"),
        ("Bahasa interface", "Bahasa Indonesia (multi-bahasa direncanakan)"),
        ("Platform target", "Android 7.0+, iOS 12+, Web (chromium-based)"),
        ("Kapasitas user", "Single-user per device; tidak ada multi-akun lokal"),
        ("Ukuran model lokal", "~1.2 GB (HPT-D) + ~360 MB (anas classifier)"),
    ],
    col_widths=[1.8, 4.6],
)

page_break(doc)

# ============ 5. ARSITEKTUR SISTEM ============
heading(doc, "5. Arsitektur Sistem End-to-End", level=1)

heading(doc, "5.1 Diagram Arsitektur Konseptual", level=2)
para(doc, "Arsitektur sistem terdiri dari tiga layer utama:")

code_block(doc, """\
+-----------------------------------------------------------+
|                    MOBILE CLIENT (Flutter)                |
|  Splash → Auth → Home → Jilid → Lesson → Eval → Progress  |
|  ┌─────────────┐  ┌──────────────┐  ┌───────────────┐    |
|  │  Provider   │  │  ApiService  │  │  AudioRecorder │    |
|  │  (state)    │  │  (HTTP+JWT)  │  │  (16kHz WAV)   │    |
|  └─────────────┘  └──────────────┘  └───────────────┘    |
+-----------------------------------------------------------+
              │ HTTPS (multipart/form-data)
              ▼
+-----------------------------------------------------------+
|                   BACKEND (FastAPI Async)                 |
|  Routers ─→ Services ─→ Models (SQLAlchemy)               |
|  ┌─────────────┐  ┌──────────────┐  ┌───────────────┐    |
|  │  auth.py    │  │ ai_service.py│  │  classifier.py │    |
|  │  evaluation │  │ (HPT-D wrap) │  │  (anas wrap)   │    |
|  │  lessons    │  │  asr_inference│  │                │    |
|  │  progress   │  │              │  │                │    |
|  │  glossary   │  │              │  │                │    |
|  └─────────────┘  └──────────────┘  └───────────────┘    |
+-----------------------------------------------------------+
       │                       │                  │
       ▼                       ▼                  ▼
+-------------+    +--------------------+    +----------------+
| PostgreSQL  |    |  Model HPT-D       |    |  Model anas    |
| users,      |    |  (Wav2Vec2ForCTC)  |    |  (Classifier)  |
| evals,      |    |  CER 0.10, WER 0.20|    |  84.28% test   |
| progress    |    |  TSA 65.2%         |    |  acc           |
+-------------+    +--------------------+    +----------------+
""")

heading(doc, "5.2 Komponen Utama", level=2)
make_table(doc,
    ["Layer", "Komponen", "Tanggung Jawab"],
    [
        ("Mobile", "Flutter App", "UI, audio capture, state management, persistensi token"),
        ("Mobile", "Provider", "ChangeNotifier untuk auth, lesson, evaluation, progress"),
        ("Mobile", "ApiService", "HTTP client + JWT injection + multipart upload"),
        ("Backend", "FastAPI", "Async API gateway, routing, DI"),
        ("Backend", "Auth Service", "Password hashing (bcrypt), JWT generation/validation"),
        ("Backend", "AI Service", "Wrapper untuk HPT-D (ASR) dan anas (classification)"),
        ("Backend", "Lessons Service", "Konten Tilawati Jilid 1-6 (in-memory dict)"),
        ("Storage", "PostgreSQL", "Data user, evaluasi, progres (async via asyncpg)"),
        ("Storage", "Filesystem", "Audio uploads (UPLOAD_DIR)"),
        ("Model", "HPT-D", "ASR untuk rangkaian 2-4 huruf hijaiyah (Wav2Vec2ForCTC)"),
        ("Model", "anas classifier", "Klasifikasi 1 huruf hijaiyah (Wav2Vec2ForSeqCls)"),
    ],
    col_widths=[1.0, 1.6, 3.8],
)

heading(doc, "5.3 Aliran Data Latihan Per-Kata", level=2)
numbered(doc, "User membuka halaman pelajaran (Jilid X, Halaman Y).")
numbered(doc, "Flutter mem-parse transliterasi & teks Arab menjadi list grup kata.")
numbered(doc, "Grup pertama disorot hijau dengan underline (active state).")
numbered(doc, "User tap mikrofon → record package mulai capture audio 16kHz mono WAV.")
numbered(doc, "User mengucapkan grup, lalu tap stop → audio disimpan ke temp directory.")
numbered(doc, "Flutter POST multipart ke /api/evaluation/submit_word dengan field `expected_word` dan `audio`.")
numbered(doc, "Backend menyimpan audio ke temp file, panggil transcribe_word() (HPT-D inference).")
numbered(doc, "Backend membandingkan expected sequence vs ASR output sequence → matched true/false.")
numbered(doc, "Response dikirim balik ke Flutter dalam bentuk JSON {matched, asr_transcript, expected}.")
numbered(doc, "Flutter update state grup: hijau jika matched, merah jika tidak. Pindah ke grup berikutnya.")
numbered(doc, "Saat semua grup selesai → tampilkan ringkasan skor dengan grade Mumtaz/Jayyid Jiddan/Jayyid/Maqbul.")

heading(doc, "5.4 Aliran Data Glosarium", level=2)
numbered(doc, "User membuka Glosarium dari Home Screen.")
numbered(doc, "Flutter GET /api/glossary/letters → menerima list 28 huruf dengan metadata.")
numbered(doc, "User tap salah satu huruf → masuk ke detail screen.")
numbered(doc, "User tap mikrofon, ucapkan huruf, stop.")
numbered(doc, "Flutter POST multipart ke /api/glossary/classify dengan audio.")
numbered(doc, "Backend memanggil classification_inference (model anas) → return huruf prediksi + confidence.")
numbered(doc, "Flutter menampilkan huruf prediksi + confidence dan apakah cocok dengan huruf yang sedang dipelajari.")

heading(doc, "5.5 Tech Stack", level=2)
make_table(doc,
    ["Layer", "Teknologi", "Versi"],
    [
        ("Mobile Framework", "Flutter (Dart)", "3.x"),
        ("Mobile State", "provider", "^6.1.2"),
        ("Mobile HTTP", "http", "^1.2.2"),
        ("Mobile Audio Rec", "record", "any"),
        ("Mobile Storage", "shared_preferences", "^2.3.2"),
        ("Mobile UI", "google_fonts, percent_indicator, fl_chart, lottie", "—"),
        ("Backend", "FastAPI", "0.115+"),
        ("Backend ORM", "SQLAlchemy 2 Async", "—"),
        ("Backend Auth", "python-jose, passlib[bcrypt]", "—"),
        ("Database", "PostgreSQL", "15+ via asyncpg"),
        ("ML Framework", "Transformers (HuggingFace)", "4.46+ (lokal), 4.45 (training)"),
        ("ML Backbone", "Wav2Vec2 XLS-R Arabic (anas)", "—"),
        ("Audio Process", "librosa, numpy, soundfile", "—"),
        ("Container", "Docker + docker-compose", "—"),
    ],
    col_widths=[1.5, 3.0, 1.2],
)

page_break(doc)

# ============ 6. KOMPONEN AI/ML ============
heading(doc, "6. Komponen AI/ML — Model HPT-D & Klasifikasi", level=1)

heading(doc, "6.1 Pipeline Pengembangan Model", level=2)
para(doc,
    "Model dikembangkan dalam dua fase R&D yang berurutan: Klasifikasi (Phase 1) dan ASR (Phase 2). "
    "Phase 1 menghasilkan baseline yang juga di-reuse sebagai backbone untuk Phase 2."
)

heading(doc, "Phase 1 — Klasifikasi 28 Huruf", level=3)
bullet(doc, "Task: klasifikasi 1 huruf per rekaman (single-class)")
bullet(doc, "Arsitektur: Wav2Vec2ForSequenceClassification + CrossEntropy + Label Smoothing")
bullet(doc, "Backbone: anas/wav2vec2-large-xlsr-arabic (frozen feature extractor)")
bullet(doc, "Hasil: 84.28% test accuracy")
bullet(doc, "Penggunaan saat ini: glosarium per-huruf di TilawatiMobile")

heading(doc, "Phase 2 — ASR Sequence (HPT-D)", level=3)
bullet(doc, "Task: transkripsi rangkaian 2-4 huruf per rekaman")
bullet(doc, "Arsitektur: Wav2Vec2ForCTC + CTC Loss")
bullet(doc, "Vocab: word-level dengan 31 entries (huruf hijaiyah sebagai 'kata')")
bullet(doc, "Backbone: fine-tuned dari hasil Phase 1 (transfer learning)")
bullet(doc, "Augmentasi: Gaussian Noise + Time Shift + Pitch Shift, hanya saat collator (bukan dataset.map)")
bullet(doc, "Oversampling: TSA × 1.5 (untuk mencegah collapse pada huruf TSA)")
bullet(doc, "Hasil final HPT-D: CER 0.1013, WER 0.2025, TSA per-class accuracy 65.2%")

heading(doc, "6.2 Spesifikasi Model HPT-D", level=2)
make_table(doc,
    ["Aspek", "Detail"],
    [
        ("Tipe Arsitektur", "Wav2Vec2ForCTC (HuggingFace Transformers)"),
        ("Backbone", "anas/wav2vec2-large-xlsr-arabic (fine-tuned)"),
        ("Output Vocab", "Word-level, 31 entries (28 huruf + 3 token kontrol)"),
        ("Tokenizer Size", "33 (dengan BOS/EOS tambahan)"),
        ("Sample Rate Input", "16 kHz mono"),
        ("Format Input", "WAV float32"),
        ("Decode Strategy", "Manual CTC collapse via ID2WORD (bukan tokenizer.batch_decode)"),
        ("Ukuran Model", "~1.2 GB (full precision)"),
        ("CER (Character Error Rate)", "0.1013 (10.13%)"),
        ("WER (Word Error Rate)", "0.2025 (20.25%)"),
        ("TSA per-class accuracy", "65.2% (huruf paling sulit)"),
        ("Inference Latency (Local)", "~500 ms per audio 2-3 detik (CPU)"),
    ],
    col_widths=[2.0, 4.4],
)

heading(doc, "6.3 Inference Pipeline Backend", level=2)
para(doc, "Modul backend bertanggung jawab untuk inference:")
bullet(doc, "backend/app/services/asr_inference.py — load_asr_model() + transcribe_file()")
bullet(doc, "backend/app/services/classification_inference.py — load_classification_model() + classify_file()")
bullet(doc, "backend/app/services/ai_service.py — public wrapper: evaluate_audio, transcribe_word, classify_audio")

para(doc, "Loading Strategy:")
bullet(doc, "Model di-load saat lifespan startup (sekali per worker), tidak per request.")
bullet(doc, "load_asr_model() dipanggil dalam executor (loop.run_in_executor) untuk tidak memblok event loop.")
bullet(doc, "Path model dikontrol via ENV: ASR_MODEL_DIR dan CLASSIFICATION_MODEL_DIR.")
bullet(doc, "from_pretrained() dipanggil dengan local_files_only=True untuk Transformers ≥ 4.46.")

para(doc, "Fallback Strategy:")
bullet(doc, "Jika ASR_SERVICE_URL diset (HuggingFace Spaces), backend bertindak sebagai proxy.")
bullet(doc, "Jika model lokal tidak tersedia DAN URL remote tidak diset, evaluasi fallback ke mock 70% match rate.")
bullet(doc, "Fallback memastikan demo tetap berjalan saat infrastruktur model belum siap.")

heading(doc, "6.4 Constraint Audio Input", level=2)
make_table(doc,
    ["Parameter", "Nilai", "Alasan"],
    [
        ("Sample Rate", "16 kHz", "Match training data; Wav2Vec2 standar 16 kHz"),
        ("Channels", "Mono", "Klasifikasi & ASR tidak butuh stereo; menghemat bandwidth"),
        ("Format", "WAV (PCM)", "Lossless, kompatibel dengan record + librosa"),
        ("Durasi minimum", "0.3 detik", "Di bawah ini terlalu pendek untuk frame 25ms"),
        ("Durasi maksimum", "10 detik", "Di atas ini biasanya bukan 1-4 huruf, model mis-decode"),
        ("Encoding bit depth", "16-bit signed", "Default record package; cukup untuk SNR tinggi"),
    ],
    col_widths=[1.4, 1.2, 3.8],
)

heading(doc, "6.5 Glosarium dengan Model Klasifikasi anas", level=2)
para(doc,
    "Untuk fitur glosarium per-huruf, dipakai model klasifikasi (Phase 1) dengan task single-class. "
    "Endpoint /api/glossary/classify menerima audio satu huruf, mengembalikan huruf prediksi + "
    "confidence score. Model ini lebih akurat untuk single-class (84.28% vs 65.2% TSA HPT-D) sehingga "
    "lebih cocok untuk drill huruf demi huruf."
)

page_break(doc)

# ============ 7. ARSITEKTUR BACKEND ============
heading(doc, "7. Arsitektur Backend (FastAPI)", level=1)

heading(doc, "7.1 Module Structure", level=2)
code_block(doc, """\
backend/
├── app/
│   ├── main.py              # FastAPI instance, lifespan, CORS, router registration
│   ├── config.py            # Settings dari ENV via python-dotenv
│   ├── database.py          # SQLAlchemy async engine + get_db
│   ├── models/              # SQLAlchemy ORM (DeclarativeBase)
│   │   ├── user.py          # tabel users
│   │   ├── evaluation.py    # tabel evaluations (with JSONB feedback)
│   │   └── progress.py      # tabel progress
│   ├── schemas/             # Pydantic models (request/response)
│   │   ├── user.py
│   │   ├── evaluation.py
│   │   └── progress.py
│   ├── routers/             # API endpoints
│   │   ├── auth.py          # /api/auth/*
│   │   ├── lessons.py       # /api/lessons/*
│   │   ├── evaluation.py    # /api/evaluation/*
│   │   ├── progress.py      # /api/progress/*
│   │   └── glossary.py      # /api/glossary/*
│   ├── services/            # Business logic
│   │   ├── auth_service.py  # JWT + password hashing
│   │   ├── ai_service.py    # public AI wrappers
│   │   ├── asr_inference.py # HPT-D loader + transcriber
│   │   └── classification_inference.py
│   └── utils/
│       └── audio.py
├── requirements.txt
└── .env.example
""")

heading(doc, "7.2 Application Lifecycle", level=2)
para(doc,
    "FastAPI application memakai async context manager (lifespan) untuk inisialisasi resource startup "
    "dan cleanup saat shutdown. Detail di app/main.py:"
)
bullet(doc, "Buat UPLOAD_DIR jika belum ada.")
bullet(doc, "Inisialisasi database: jalankan create_all (idempotent) untuk auto-create tabel.")
bullet(doc, "Load model HPT-D dan classifier dalam thread executor (non-blocking).")
bullet(doc, "Print readiness signal (✅ Tilawati API is ready!).")
bullet(doc, "Saat shutdown: print farewell.")

heading(doc, "7.3 CORS & Middleware", level=2)
para(doc,
    "CORS middleware dikonfigurasi sangat permisif untuk fase development: allow_origins=['*'], "
    "allow_methods=['*'], allow_headers=['*']. Untuk produksi, harus diperketat ke domain Flutter "
    "deployment (misal https://tilawati.app) atau set token-based origin validation."
)

heading(doc, "7.4 Autentikasi JWT", level=2)
para(doc, "Implementasi auth_service.py:")
bullet(doc, "Password di-hash dengan bcrypt (passlib).")
bullet(doc, "JWT signed dengan HS256 + SECRET_KEY dari ENV.")
bullet(doc, "Access token expires dalam ACCESS_TOKEN_EXPIRE_MINUTES (default 1440 menit = 24 jam).")
bullet(doc, "Dependency get_current_user() decode token dari Authorization header (Bearer xxx).")
bullet(doc, "Endpoint terlindungi dengan Depends(get_current_user) — non-authenticated request mendapat 401.")

heading(doc, "7.5 File Handling Audio", level=2)
para(doc,
    "Strategi file handling berbeda per endpoint berdasarkan kebutuhan persistensi:"
)
bullet(doc, "/api/evaluation/submit (full lesson): audio disimpan permanen ke UPLOAD_DIR dengan format {user_id}_{jilid}_{lesson}_{hash}.wav untuk audit.")
bullet(doc, "/api/evaluation/submit_word (per kata): audio disimpan sementara via tempfile.NamedTemporaryFile, dihapus segera setelah inference.")
bullet(doc, "/api/glossary/classify: pola sama dengan submit_word (temporary).")

heading(doc, "7.6 Database Layer", level=2)
para(doc, "Konfigurasi database/SQLAlchemy:")
bullet(doc, "Engine async via create_async_engine + asyncpg driver.")
bullet(doc, "DATABASE_URL: postgresql+asyncpg://tilawati:tilawati123@localhost:5432/tilawati_db (default dev).")
bullet(doc, "Session factory: async_sessionmaker(bind=engine, expire_on_commit=False).")
bullet(doc, "DI dependency get_db() yields AsyncSession dengan auto-rollback on exception.")
bullet(doc, "Schema dibuat via Base.metadata.create_all() di startup (DEV); produksi pakai Alembic migrations.")

heading(doc, "7.7 Service Layer (ai_service.py)", level=2)
para(doc, "ai_service.py mengekspos tiga fungsi public:")
bullet(doc, "evaluate_audio(audio_path, jilid, lesson_number) → evaluasi seluruh halaman, return skor + breakdown")
bullet(doc, "transcribe_word(audio_path) → list[str] hasil ASR (untuk submit_word)")
bullet(doc, "classify_audio(audio_path) → (huruf, confidence) untuk glosarium")
bullet(doc, "get_lesson(jilid, lesson_number), get_jilid_lessons(jilid) — akses konten Tilawati statik")

para(doc, "Konten Tilawati disimpan in-memory sebagai Python dict pada module load. Untuk skala besar atau update konten dinamis, ke depan dapat dipindah ke tabel `lessons` di PostgreSQL.")

heading(doc, "7.8 Error Handling", level=2)
make_table(doc,
    ["Status Code", "Konteks", "Pesan"],
    [
        ("400 Bad Request", "Validasi input (Pydantic)", "Field required atau invalid"),
        ("401 Unauthorized", "Token JWT invalid/expired", "Token expired or invalid"),
        ("404 Not Found", "Lesson/evaluation tidak ditemukan", "Pelajaran tidak ditemukan"),
        ("422 Unprocessable", "Form/multipart parsing error", "—"),
        ("500 Internal", "ASR/DB error", "Error di-log via uvicorn"),
    ],
    col_widths=[1.4, 2.0, 3.0],
)

page_break(doc)

# ============ 8. ARSITEKTUR FRONTEND ============
heading(doc, "8. Arsitektur Frontend (Flutter)", level=1)

heading(doc, "8.1 Module Structure", level=2)
code_block(doc, """\
frontend/
└── lib/
    ├── main.dart                  # MultiProvider + MaterialApp + routes
    ├── config/
    │   ├── theme.dart             # AppColors, AppTheme
    │   └── api_config.dart        # Base URL + semua endpoint URL
    ├── models/
    │   ├── user_model.dart
    │   ├── jilid_model.dart
    │   └── evaluation_model.dart
    ├── services/
    │   └── api_service.dart       # HTTP wrapper (get/post/postMultipart) + JWT
    ├── providers/
    │   ├── auth_provider.dart
    │   ├── lesson_provider.dart
    │   ├── evaluation_provider.dart
    │   └── progress_provider.dart
    └── screens/
        ├── splash_screen.dart
        ├── login_screen.dart
        ├── register_screen.dart
        ├── home_screen.dart
        ├── jilid_selection_screen.dart   # konten lokal Jilid + lesson list
        ├── lesson_screen.dart            # rekam per-kata, highlight Arabic
        ├── evaluation_screen.dart        # hasil evaluasi (full mode)
        ├── progress_screen.dart          # dashboard
        ├── history_screen.dart           # riwayat
        ├── glossary_screen.dart          # 28 huruf
        └── glossary_detail_screen.dart
""")

heading(doc, "8.2 State Management — Provider", level=2)
para(doc, "Aplikasi memakai pattern provider (ChangeNotifier) untuk state global. Empat provider utama didaftarkan di MultiProvider pada main.dart:")
bullet(doc, "AuthProvider — login state, current user, JWT token persistence.")
bullet(doc, "LessonProvider — daftar jilid, detail lesson dari /api/lessons (sebagian besar konten lokal).")
bullet(doc, "EvaluationProvider — submit evaluasi, fetch history, expose model + ASR transcript.")
bullet(doc, "ProgressProvider — fetch dashboard, expose totalEvaluations, averageScore, jilidProgress, recentEvaluations.")

heading(doc, "8.3 Service Layer — ApiService", level=2)
para(doc, "ApiService (singleton-like static class) menyediakan:")
bullet(doc, "loadToken() / saveToken() / clearToken() — persistence via SharedPreferences.")
bullet(doc, "get(url) — auto-inject Authorization: Bearer header jika token tersedia.")
bullet(doc, "post(url, body) — JSON post dengan token.")
bullet(doc, "postMultipart(url, fields, filePath, fileField) — upload audio dengan token.")

heading(doc, "8.4 Audio Recording", level=2)
para(doc, "Implementasi menggunakan package `record`:")
code_block(doc, """\
await _recorder.start(
  const RecordConfig(
    encoder: AudioEncoder.wav,
    sampleRate: 16000,
    numChannels: 1,
  ),
  path: '${dir.path}/tilawati_word_${timestamp}.wav',
);""")
bullet(doc, "Path file di-resolve via path_provider.getTemporaryDirectory().")
bullet(doc, "Permission dicek via permission_handler sebelum record.start.")
bullet(doc, "Timer durasi rekaman ditampilkan ke user (mm:ss).")
bullet(doc, "Pulse animation pada tombol mic untuk visual feedback.")

heading(doc, "8.5 Routing & Navigation", level=2)
para(doc, "Routes terdaftar di MaterialApp.routes:")
make_table(doc,
    ["Path", "Screen", "Argumen"],
    [
        ("/", "SplashScreen", "—"),
        ("/login", "LoginScreen", "—"),
        ("/register", "RegisterScreen", "—"),
        ("/home", "HomeScreen", "—"),
        ("/jilid-selection", "JilidSelectionScreen", "—"),
        ("/lesson", "LessonScreen", "{jilid, lesson, color}"),
        ("/evaluation", "EvaluationScreen", "{jilid, lesson, audioPath, color}"),
        ("/progress", "ProgressScreen", "—"),
        ("/glossary", "GlossaryScreen", "—"),
        ("/glossary-detail", "GlossaryDetailScreen", "{letter}"),
        ("/riwayat", "HistoryScreen", "—"),
    ],
    col_widths=[1.5, 2.0, 3.0],
)

heading(doc, "8.6 Theming", level=2)
para(doc, "Tema didefinisikan di config/theme.dart:")
bullet(doc, "AppColors.primary: hijau Tilawati (0xFF1B5E20)")
bullet(doc, "Tipografi: Google Fonts Poppins (sans-serif, ramah anak)")
bullet(doc, "Material Design 3 dengan Color Scheme custom")
bullet(doc, "Per-jilid color: setiap jilid punya warna khas (Jilid 1 hijau, 2 biru, 3 ungu, dst.) untuk identitas visual")

heading(doc, "8.7 LessonScreen — Komponen Inti Latihan", level=2)
para(doc, "LessonScreen adalah komponen paling kompleks. Tanggung jawabnya:")
bullet(doc, "Parse transliterasi & teks Arab menjadi list grup (per 'kata'/grup huruf yang dipisahkan 2+ spasi).")
bullet(doc, "Render Arabic text dengan per-grup styling: pending (default), active (hijau + underline border-bottom), correct (hijau muda), incorrect (merah).")
bullet(doc, "Mengelola state machine recording: idle → recording → evaluating → next group.")
bullet(doc, "Setelah semua grup selesai → tampilkan _buildScoreSummary() dengan grade.")
para(doc, "Catatan teknis terbaru (30 April 2026):")
bullet(doc, "Underline grup aktif memakai Container border-bottom dengan padding bottom: 2 px (bukan TextDecoration.underline) agar tidak memotong harakat fathah pada huruf Arab.")
bullet(doc, "Transliterasi tidak ditampilkan lagi di lesson screen — pengguna fokus pada teks Arab.")
bullet(doc, "Token grouping memakai split RegExp(r'  +') (2+ spasi) sebagai delimiter antar kata.")

page_break(doc)

# ============ 9. SPESIFIKASI FUNGSIONAL ============
heading(doc, "9. Spesifikasi Fungsional", level=1)

heading(doc, "9.1 User Stories", level=2)
make_table(doc,
    ["ID", "Sebagai", "Saya ingin", "Sehingga"],
    [
        ("US-01", "Pemula", "Mendaftar dengan email dan password", "Bisa menyimpan progres saya"),
        ("US-02", "User terdaftar", "Login dengan kredensial saya", "Akses ke fitur lengkap"),
        ("US-03", "User", "Memilih jilid yang sesuai level saya", "Tidak overload dengan materi terlalu sulit"),
        ("US-04", "User", "Melihat halaman pelajaran (Arab + transliterasi)", "Tahu bacaan yang harus diucapkan"),
        ("US-05", "User", "Merekam bacaan per kata / grup huruf", "Mendapat feedback bertahap"),
        ("US-06", "User", "Melihat indikator hijau/merah per grup", "Tahu mana yang benar/salah segera"),
        ("US-07", "User", "Melihat skor akhir + grade per pelajaran", "Mengukur progres keseluruhan"),
        ("US-08", "User", "Mengulang pelajaran yang skor rendah", "Memperbaiki kelemahan"),
        ("US-09", "User", "Membuka glosarium 28 huruf", "Melatih huruf demi huruf"),
        ("US-10", "User", "Melihat riwayat evaluasi", "Memantau perkembangan kemampuan"),
        ("US-11", "User", "Melihat dashboard progres per jilid", "Tahu seberapa jauh saya sudah berkembang"),
    ],
    col_widths=[0.7, 1.2, 2.4, 2.0],
)

heading(doc, "9.2 Use Case Utama — Latihan Per Kata", level=2)
make_table(doc,
    ["Aspek", "Detail"],
    [
        ("Aktor Utama", "User terdaftar"),
        ("Pre-condition", "User sudah login dengan token valid; backend dapat diakses; mikrofon diizinkan"),
        ("Trigger", "User memilih lesson dari Jilid Selection Screen"),
        ("Main Flow",
            "1. Sistem menampilkan teks Arab dengan grup pertama disorot hijau.\n"
            "2. User tap mikrofon → recording dimulai.\n"
            "3. User mengucapkan grup huruf, lalu tap stop.\n"
            "4. Sistem upload ke /api/evaluation/submit_word.\n"
            "5. Backend menjalankan HPT-D inference, kembalikan matched true/false.\n"
            "6. Sistem update warna grup: hijau (correct) atau merah (incorrect).\n"
            "7. Setelah delay 700ms, grup berikutnya menjadi active (hijau + underline).\n"
            "8. Ulangi 2-7 sampai semua grup selesai.\n"
            "9. Sistem tampilkan _buildScoreSummary() dengan grade."),
        ("Alternative Flow",
            "5a. ASR tidak tersedia → fallback mock 70% match rate.\n"
            "5b. Network error → token ditandai incorrect, lanjut ke berikutnya."),
        ("Post-condition", "Progres user di-update; user bisa retry atau exit"),
    ],
    col_widths=[1.6, 4.8],
)

heading(doc, "9.3 Skema Penilaian Lesson Penuh", level=2)
para(doc, "Untuk endpoint /api/evaluation/submit (mode full lesson, belum aktif sebagai default UX), sistem menghitung tiga komponen skor:")
bullet(doc, "Makharijul Huruf — bobot 40% — akurasi titik artikulasi huruf (HPT-D output match)")
bullet(doc, "Tajwid — bobot 35% — saat ini placeholder, ke depan dapat disambung ke model tajwid")
bullet(doc, "Kelancaran — bobot 25% — durasi audio relatif terhadap jumlah huruf")
para(doc, "Grade akhir berdasarkan overall_score:")
make_table(doc,
    ["Skor", "Grade"],
    [
        ("≥ 85", "Mumtaz (Sangat Baik)"),
        ("≥ 70", "Jayyid Jiddan (Baik Sekali)"),
        ("≥ 55", "Jayyid (Baik)"),
        ("< 55", "Maqbul (Cukup)"),
    ],
    col_widths=[1.0, 3.0],
)

page_break(doc)

# ============ 10. NON-FUNGSIONAL ============
heading(doc, "10. Spesifikasi Non-Fungsional", level=1)

heading(doc, "10.1 Performance", level=2)
make_table(doc,
    ["Metrik", "Target", "Pengukuran"],
    [
        ("Latensi evaluasi per kata", "≤ 3 detik (LAN)", "Manual stopwatch dari tap stop ke feedback"),
        ("Latensi evaluasi per kata", "≤ 5 detik (4G)", "Sama, dengan device aktif di 4G"),
        ("Waktu boot backend", "≤ 30 detik", "Termasuk load 2 model"),
        ("Waktu start aplikasi", "≤ 3 detik", "Dari tap icon ke splash"),
        ("Memory footprint Flutter", "≤ 200 MB", "Profile mode di Android"),
        ("Memory footprint Backend", "≤ 4 GB", "Dengan HPT-D dan classifier loaded"),
    ],
    col_widths=[2.0, 1.6, 2.8],
)

heading(doc, "10.2 Security", level=2)
bullet(doc, "Password disimpan sebagai bcrypt hash, bukan plaintext.")
bullet(doc, "JWT signed dengan SECRET_KEY (harus diganti di produksi, jangan hardcode).")
bullet(doc, "JWT expiry 24 jam (default), refresh harus login ulang.")
bullet(doc, "CORS terbuka di dev — wajib diperketat di produksi (origin whitelist).")
bullet(doc, "Audio upload size diberi limit (FastAPI default ~16 MB; sesuaikan dengan kebutuhan).")
bullet(doc, "User_id divalidasi via JWT untuk semua endpoint yang protected.")
bullet(doc, "Cascade delete: hapus user → otomatis hapus evaluations & progress (foreign key ON DELETE CASCADE).")

heading(doc, "10.3 Usability", level=2)
bullet(doc, "Material Design 3 dengan tipografi Poppins yang ramah anak.")
bullet(doc, "Visual feedback dominan warna (hijau/merah) tanpa bergantung teks panjang — cocok untuk anak yang belum lancar membaca Indonesia.")
bullet(doc, "Tombol mic besar (≥ 80px) dengan animasi pulse.")
bullet(doc, "Layout responsive untuk berbagai ukuran layar (handphone & tablet).")
bullet(doc, "Tipografi Arab khusus (font 'Arial' fallback dengan height 2 untuk harakat tidak terpotong).")

heading(doc, "10.4 Reliability", level=2)
bullet(doc, "Fallback mock 70% match rate jika ASR tidak tersedia → demo tetap jalan.")
bullet(doc, "Progress disimpan di backend → tidak hilang saat reinstall app.")
bullet(doc, "Token JWT di-cache via SharedPreferences → user tidak perlu login berulang.")
bullet(doc, "Database transaction di-rollback otomatis on exception.")
bullet(doc, "Async + threadpool executor mencegah model loading memblok event loop.")

heading(doc, "10.5 Compatibility", level=2)
bullet(doc, "Android 7.0 (API 24) ke atas — covers ~95% device aktif di Indonesia.")
bullet(doc, "iOS 12+ — minimum supported by Flutter stable.")
bullet(doc, "Web (sekunder) — Chromium-based browser, butuh getUserMedia untuk recording.")
bullet(doc, "Backend cross-platform Python ≥3.10 (Linux/macOS/Windows).")
bullet(doc, "Database PostgreSQL 13+ (testing di 15).")

heading(doc, "10.6 Maintainability", level=2)
bullet(doc, "Project terstruktur clean architecture (models, schemas, routers, services, utils).")
bullet(doc, "Type hints lengkap di backend (mypy-friendly).")
bullet(doc, "Konvensi naming konsisten: snake_case Python, camelCase Dart.")
bullet(doc, "CLAUDE.md sebagai dokumentasi internal untuk AI-assisted development.")
bullet(doc, "Provider pattern di Flutter untuk testability tinggi.")

page_break(doc)

# ============ 11. DATA MODEL ============
heading(doc, "11. Data Model & Database Schema", level=1)

heading(doc, "11.1 Skema Tabel users", level=2)
make_table(doc,
    ["Kolom", "Tipe", "Constraint", "Deskripsi"],
    [
        ("id", "UUID", "PK, default uuid4", "Primary key"),
        ("name", "VARCHAR(100)", "NOT NULL", "Nama lengkap"),
        ("email", "VARCHAR(255)", "UNIQUE, NOT NULL, INDEX", "Untuk login"),
        ("hashed_password", "VARCHAR(255)", "NOT NULL", "Bcrypt hash"),
        ("avatar_url", "VARCHAR(500)", "NULLABLE", "URL avatar (opsional)"),
        ("created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Tanggal registrasi"),
        ("updated_at", "TIMESTAMPTZ", "DEFAULT NOW(), ON UPDATE", "Update terakhir"),
    ],
    col_widths=[1.4, 1.5, 1.6, 1.9],
)

heading(doc, "11.2 Skema Tabel evaluations", level=2)
make_table(doc,
    ["Kolom", "Tipe", "Constraint", "Deskripsi"],
    [
        ("id", "UUID", "PK, default uuid4", "Primary key"),
        ("user_id", "UUID", "FK→users.id, CASCADE, INDEX", "Pemilik evaluasi"),
        ("jilid", "INT", "NOT NULL", "Jilid 1-6"),
        ("lesson_number", "INT", "NOT NULL", "Halaman dalam jilid"),
        ("lesson_title", "VARCHAR(200)", "NOT NULL", "Judul lesson untuk display"),
        ("audio_path", "VARCHAR(500)", "NULLABLE", "Path file di UPLOAD_DIR"),
        ("overall_score", "FLOAT", "DEFAULT 0", "Skor keseluruhan"),
        ("makharijul_huruf_score", "FLOAT", "DEFAULT 0", "Skor 40% bobot"),
        ("tajwid_score", "FLOAT", "DEFAULT 0", "Skor 35% bobot"),
        ("kelancaran_score", "FLOAT", "DEFAULT 0", "Skor 25% bobot"),
        ("feedback_json", "JSONB", "NULLABLE", "Grade, message, tips"),
        ("phoneme_details", "JSONB", "NULLABLE", "asr_transcript, expected, matched_tokens"),
        ("created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Waktu submit"),
    ],
    col_widths=[1.5, 1.2, 1.7, 2.0],
)

heading(doc, "11.3 Skema Tabel progress", level=2)
make_table(doc,
    ["Kolom", "Tipe", "Constraint", "Deskripsi"],
    [
        ("id", "UUID", "PK, default uuid4", "Primary key"),
        ("user_id", "UUID", "FK→users.id, CASCADE, INDEX", "Pemilik progres"),
        ("jilid", "INT", "NOT NULL", "Jilid 1-6"),
        ("total_lessons", "INT", "DEFAULT 0", "Total lesson dalam jilid"),
        ("completed_lessons", "INT", "DEFAULT 0", "Lesson unik yang sudah dievaluasi"),
        ("average_score", "FLOAT", "DEFAULT 0", "Rata-rata skor"),
        ("best_score", "FLOAT", "DEFAULT 0", "Skor tertinggi"),
        ("total_practice_count", "INT", "DEFAULT 0", "Jumlah submit (termasuk repeats)"),
        ("updated_at", "TIMESTAMPTZ", "DEFAULT NOW(), ON UPDATE", "Update terakhir"),
    ],
    col_widths=[1.5, 1.2, 1.7, 2.0],
)

heading(doc, "11.4 Relasi Antar Tabel", level=2)
code_block(doc, """\
users (1) ────┬───── (N) evaluations
              │       (CASCADE on delete)
              │
              └───── (N) progress
                      (CASCADE on delete; UNIQUE per user_id+jilid logically)
""")
bullet(doc, "Setiap user memiliki banyak evaluations (N evaluasi historis per user).")
bullet(doc, "Setiap user memiliki maksimal 6 progress records (1 per jilid).")
bullet(doc, "Cascade delete: menghapus user → otomatis hapus semua evaluations dan progress miliknya.")

heading(doc, "11.5 Indexing Strategy", level=2)
bullet(doc, "users.email — UNIQUE INDEX (untuk lookup login cepat).")
bullet(doc, "evaluations.user_id — INDEX (untuk query history per user).")
bullet(doc, "progress.user_id — INDEX (untuk query dashboard per user).")
bullet(doc, "evaluations(user_id, jilid, created_at DESC) — composite index ke depan untuk query history dengan filter jilid.")

page_break(doc)

# ============ 12. API SPEC ============
heading(doc, "12. API Specification", level=1)

heading(doc, "12.1 Konvensi", level=2)
bullet(doc, "Base URL: http://10.0.2.2:8000 (Android emulator) / http://localhost:8000 (iOS) / http://<LAN-IP>:8000 (real device)")
bullet(doc, "Prefix API: /api")
bullet(doc, "Format request body: JSON, kecuali endpoint upload (multipart/form-data)")
bullet(doc, "Format response: JSON")
bullet(doc, "Authentication: Header `Authorization: Bearer <token>` untuk endpoint protected")

heading(doc, "12.2 Authentication Endpoints", level=2)
make_table(doc,
    ["Method", "Path", "Auth", "Deskripsi"],
    [
        ("POST", "/api/auth/register", "Public", "Daftar user baru. Body: {name, email, password}"),
        ("POST", "/api/auth/login", "Public", "Login. Body: {email, password}. Return: {access_token, user}"),
        ("GET", "/api/auth/me", "Required", "Info user saat ini"),
    ],
    col_widths=[0.8, 2.2, 1.0, 2.4],
)

heading(doc, "12.3 Lessons Endpoints", level=2)
make_table(doc,
    ["Method", "Path", "Auth", "Deskripsi"],
    [
        ("GET", "/api/lessons/jilids", "Required", "List 6 jilid dengan metadata"),
        ("GET", "/api/lessons/jilids/{id}", "Required", "Detail jilid + list halaman"),
        ("GET", "/api/lessons/jilids/{id}/lessons/{n}", "Required", "Detail satu halaman (Arab + transliterasi)"),
    ],
    col_widths=[0.8, 3.0, 1.0, 1.6],
)

heading(doc, "12.4 Evaluation Endpoints", level=2)
make_table(doc,
    ["Method", "Path", "Auth", "Deskripsi"],
    [
        ("POST", "/api/evaluation/submit", "Required", "Submit audio satu halaman penuh"),
        ("POST", "/api/evaluation/submit_word", "Required", "Evaluasi audio per kata (instant feedback)"),
        ("GET", "/api/evaluation/history", "Required", "Riwayat evaluasi user (?jilid optional, ?limit optional)"),
        ("GET", "/api/evaluation/{id}", "Required", "Detail satu evaluasi"),
    ],
    col_widths=[0.8, 2.8, 1.0, 1.8],
)

heading(doc, "submit_word — Detail Request/Response", level=3)
para(doc, "Request (multipart/form-data):")
code_block(doc, """\
expected_word: "A BA"           # contoh grup
audio: file.wav                 # 16kHz mono WAV""")

para(doc, "Response (JSON):")
code_block(doc, """\
{
  "matched": true,
  "asr_transcript": "A BA",
  "expected": "A BA"
}""")
para(doc,
    "Logika matching: backend split expected ke tokens (['A', 'BA']), bandingkan list dengan ASR output "
    "([t.upper() for t in predicted_tokens]). matched=True jika kedua list identik (exact sequence)."
)

heading(doc, "submit — Detail Request/Response", level=3)
para(doc, "Request (multipart/form-data):")
code_block(doc, """\
jilid: 1
lesson_number: 1
audio: file.wav""")

para(doc, "Response (JSON):")
code_block(doc, """\
{
  "id": "uuid",
  "jilid": 1,
  "lesson_number": 1,
  "lesson_title": "Hal 1 - Huruf Alif & Ba",
  "overall_score": 78.5,
  "makharijul_huruf_score": 82.3,
  "tajwid_score": 74.1,
  "kelancaran_score": 79.8,
  "feedback_json": {
    "grade": "Jayyid Jiddan",
    "message": "Bacaan baik, perhatikan makharijul huruf 'A.",
    "tips": ["Ulangi latihan huruf 'A", "Rekam di tempat lebih sepi"]
  },
  "phoneme_details": {
    "asr_transcript": "ba ta tsa",
    "expected": "BA TA TSA",
    "matched_tokens": ["BA", "TA"]
  },
  "created_at": "2026-04-30T10:15:30Z"
}""")

heading(doc, "12.5 Progress Endpoints", level=2)
make_table(doc,
    ["Method", "Path", "Auth", "Deskripsi"],
    [
        ("GET", "/api/progress/dashboard", "Required", "Ringkasan progres seluruh jilid"),
        ("GET", "/api/progress/jilid/{id}", "Required", "Progres detail satu jilid"),
    ],
    col_widths=[0.8, 2.5, 1.0, 2.1],
)

heading(doc, "12.6 Glossary Endpoints", level=2)
make_table(doc,
    ["Method", "Path", "Auth", "Deskripsi"],
    [
        ("GET", "/api/glossary/letters", "Required", "List 28 huruf hijaiyah dengan metadata"),
        ("POST", "/api/glossary/classify", "Required", "Klasifikasi audio satu huruf via model anas"),
    ],
    col_widths=[0.8, 2.6, 1.0, 2.0],
)

page_break(doc)

# ============ 13. USER FLOW ============
heading(doc, "13. User Flow & Skenario Penggunaan", level=1)

heading(doc, "13.1 Onboarding Flow", level=2)
code_block(doc, """\
[Splash 2s] → [Cek token di SharedPreferences]
   ├─ Token valid    → Home Screen
   └─ Token tidak ada → Login Screen
                          ├─ "Daftar dulu" → Register Screen → Auto-login → Home
                          └─ Login → Home Screen
""")

heading(doc, "13.2 Latihan Flow", level=2)
code_block(doc, """\
Home Screen
   └─ Tap "Mulai Latihan"
        └─ Jilid Selection Screen (grid 6 jilid)
             └─ Tap Jilid X
                  └─ List 44 halaman (Jilid 1) atau 2-4 (Jilid 2-6)
                       └─ Tap Halaman Y
                            └─ Lesson Screen
                                 ├─ Tampilkan teks Arab dengan grup pertama active (hijau+underline)
                                 ├─ Loop:
                                 │    ├─ Tap mic → recording
                                 │    ├─ Tap stop → upload
                                 │    ├─ Backend → HPT-D inference
                                 │    ├─ Update grup: hijau (correct) atau merah (incorrect)
                                 │    └─ Pindah ke grup berikutnya
                                 └─ Selesai → Score Summary → [Ulangi | Selesai]
""")

heading(doc, "13.3 History Flow", level=2)
code_block(doc, """\
Home Screen → Tap "Riwayat" → History Screen
   └─ List evaluasi (terbaru di atas, 20 item)
        └─ Tap item → Detail Screen (TBD, future)
""")

heading(doc, "13.4 Progress Flow", level=2)
code_block(doc, """\
Home Screen → Tap "Progres" → Progress Screen
   ├─ Stat Cards: total latihan, rata-rata, tertinggi, streak
   ├─ Progress bars per jilid (% completed)
   └─ Aktivitas Terbaru (10 evaluasi terakhir)
""")

heading(doc, "13.5 Glosarium Flow", level=2)
code_block(doc, """\
Home Screen → Tap "Glosarium" → Glossary Screen
   └─ Grid 28 huruf hijaiyah
        └─ Tap huruf → Glossary Detail Screen
             ├─ Lihat huruf besar
             ├─ Tap mic → ucapkan huruf → stop
             ├─ Backend → classification_inference (anas)
             └─ Tampilkan: prediksi huruf + confidence + cocok/tidak
""")

page_break(doc)

# ============ 14. ROADMAP ============
heading(doc, "14. Roadmap Pengembangan", level=1)

heading(doc, "Timeline & Fase", level=2)
make_table(doc,
    ["Fase", "Periode", "Output", "Status"],
    [
        ("Fase 1: Klasifikasi", "Q3-Q4 2025", "Model klasifikasi 84.28% acc", "Selesai"),
        ("Fase 2: ASR R&D", "Q1 2026", "Model HPT-D (CER 10%, WER 20%)", "Selesai"),
        ("Fase 3: Mobile App + Backend", "Feb-Apr 2026", "TilawatiMobile MVP", "Tahap akhir"),
        ("Fase 4: Cloud Deployment", "Mei 2026", "HuggingFace Spaces, APK", "Belum"),
        ("Fase 5: Konten & Fitur Ekspansi", "Jun-Aug 2026", "Jilid 2-6 lengkap, audio ref, kuis", "Belum"),
        ("Fase 6: Pasca-TA", "Sep 2026+", "Multi-bahasa, integrasi lembaga", "Vision"),
    ],
    col_widths=[1.6, 1.2, 2.4, 1.2],
)

heading(doc, "14.1 Fase 1 — Klasifikasi (Selesai)", level=2)
bullet(doc, "Pengumpulan dataset 28 huruf hijaiyah berharakat fathah.")
bullet(doc, "Eksperimen Wav2Vec2ForSequenceClassification dengan beragam loss function.")
bullet(doc, "Final: anas/wav2vec2-large-xlsr-arabic + CrossEntropy + Label Smoothing.")
bullet(doc, "Output: 84.28% test accuracy. Model di-reuse sebagai backbone Phase 2.")

heading(doc, "14.2 Fase 2 — ASR Sequence (Selesai)", level=2)
bullet(doc, "Eksperimen vocabulary level: character-level GAGAL (collapse ke 'a', CER 0.653).")
bullet(doc, "Word-level vocabulary BERHASIL (HPT-D, CER 0.10).")
bullet(doc, "Augmentasi audio + TSA oversampling 1.5× untuk mencegah collapse pada huruf TSA.")
bullet(doc, "Hyperparameter tuning A-D, dengan HPT-D sebagai final.")

heading(doc, "14.3 Fase 3 — Mobile App + Backend (Tahap Akhir)", level=2)
para(doc, "Selesai:")
bullet(doc, "Backend FastAPI dengan auth JWT, evaluasi, progres, history, glosarium.")
bullet(doc, "Database PostgreSQL via Docker Compose.")
bullet(doc, "Model HPT-D dan classifier ter-load di startup, evaluasi end-to-end aktif.")
bullet(doc, "Konten Jilid 1 lengkap 44 halaman.")
bullet(doc, "Latihan per-kata dengan instant feedback.")
bullet(doc, "Visual highlight Arabic per grup (hijau + underline border-bottom).")
bullet(doc, "History screen + progress dashboard terhubung ke API real.")
para(doc, "Sisa pekerjaan:")
bullet(doc, "Testing E2E menyeluruh dengan emulator + device fisik.")
bullet(doc, "Polishing UI/UX (animasi transisi, error message yang lebih ramah).")
bullet(doc, "Validasi dengan partisipan demo TA.")

heading(doc, "14.4 Fase 4 — Cloud Deployment (Mei 2026)", level=2)
bullet(doc, "Deploy backend ke HuggingFace Spaces (FastAPI + model HPT-D + classifier).")
bullet(doc, "Sediakan public ASR_SERVICE_URL — Flutter dapat diuji partisipan tanpa setup lokal.")
bullet(doc, "Build APK release untuk distribusi internal (tidak via Play Store dulu).")
bullet(doc, "Setup environment ENV produksi: SECRET_KEY, DATABASE_URL pointing ke managed Postgres.")
bullet(doc, "Implementasi Alembic migrations untuk versioning skema.")

heading(doc, "14.5 Fase 5 — Konten & Fitur Ekspansi (Q3 2026)", level=2)
bullet(doc, "Konten Jilid 2 lengkap (huruf sambung).")
bullet(doc, "Konten Jilid 3-6 minimum 20 halaman per jilid.")
bullet(doc, "Audio referensi 28 huruf di glosarium.")
bullet(doc, "Mode kuis glosarium (3 mode: visual→audio, audio→visual, mixed).")
bullet(doc, "Notifikasi pengingat harian via flutter_local_notifications.")
bullet(doc, "Offline mode terbatas: cache konten + queue submit.")

heading(doc, "14.6 Fase 6 — Vision Pasca-TA", level=2)
bullet(doc, "Multi-bahasa (Arab, Indonesia, Inggris).")
bullet(doc, "Model tajwid lanjutan (Idzhar, Idgham, Ikhfa, Iqlab).")
bullet(doc, "Integrasi lembaga TPQ (multi-akun, ustadz dashboard).")
bullet(doc, "Sertifikasi level (Mumtaz Tilawati Jilid 1) yang dapat divalidasi lembaga.")
bullet(doc, "Komunitas: leaderboard, group challenges.")

page_break(doc)

# ============ 15. STATUS PROGRES ============
heading(doc, "15. Status Progres Saat Ini (30 April 2026)", level=1)

heading(doc, "15.1 Selesai", level=2)
make_table(doc,
    ["Komponen", "Detail"],
    [
        ("Auth", "Register/Login/JWT/me; password bcrypt; session via SharedPreferences"),
        ("Konten Jilid 1", "44 halaman lengkap (teks Arab + transliterasi sesuai kurikulum)"),
        ("Latihan per-kata", "Parse grup huruf, rekam, evaluasi, instant feedback"),
        ("Visual Arabic", "Per-grup styling: pending/active(hijau+underline)/correct/incorrect"),
        ("Backend submit_word", "endpoint POST /api/evaluation/submit_word, sequence matching"),
        ("Backend submit (full)", "endpoint POST /api/evaluation/submit, scoring 3 komponen"),
        ("History", "Screen + provider + endpoint /api/evaluation/history"),
        ("Progress Dashboard", "Stat cards + progress bars + recent activity"),
        ("Glosarium", "28 huruf + endpoint /api/glossary/classify (anas)"),
        ("Database", "PostgreSQL via Docker Compose"),
        ("AI Service", "Wrapper HPT-D + classifier + fallback mock"),
        ("Konten Jilid 2-6", "Minimal (2-4 lessons per jilid sebagai placeholder)"),
    ],
    col_widths=[1.8, 4.6],
)

heading(doc, "15.2 Sedang Dikerjakan / In Progress", level=2)
bullet(doc, "Polishing UX lesson screen (line tidak memotong harakat, transliterasi disembunyikan).")
bullet(doc, "Testing E2E menyeluruh — backend + frontend dengan emulator dan device real.")

heading(doc, "15.3 Belum Dimulai (Backlog)", level=2)
make_table(doc,
    ["Item", "Prioritas", "Estimasi"],
    [
        ("Konten Jilid 2-6 lengkap", "Medium", "20+ jam (per jilid)"),
        ("Audio referensi 28 huruf", "Medium", "4-6 jam (rekam + host)"),
        ("Deploy HuggingFace Spaces", "High (TA demo)", "4-6 jam"),
        ("Mode kuis glosarium", "Low", "8-10 jam"),
        ("Offline mode terbatas", "Low", "10-12 jam"),
        ("Notifikasi pengingat", "Low", "4-6 jam"),
        ("Alembic migrations", "Medium", "2-3 jam"),
        ("Rate limiting + audio validation", "Medium", "3-4 jam"),
        ("Cleanup file audio upload", "Low", "1-2 jam"),
    ],
    col_widths=[2.4, 1.6, 2.4],
)

page_break(doc)

# ============ 16. STRATEGI TESTING ============
heading(doc, "16. Strategi Testing", level=1)

heading(doc, "16.1 Unit Testing", level=2)
para(doc, "Backend (pytest):")
bullet(doc, "auth_service: hash_password, verify_password, create_access_token, decode_token.")
bullet(doc, "ai_service: parsing transliterasi, normalisasi expected vs ASR output.")
bullet(doc, "Routers: dengan TestClient + override get_db / get_current_user.")
para(doc, "Frontend (flutter_test):")
bullet(doc, "Provider state transitions (loading → success/error).")
bullet(doc, "Widget tests: token chip rendering, score summary, splash logic.")
bullet(doc, "ApiService: mock http.Client untuk verifikasi endpoint URL & header.")

heading(doc, "16.2 Integration Testing", level=2)
bullet(doc, "Backend: integration test dengan PostgreSQL test database (terpisah dari production).")
bullet(doc, "Test full flow: register → login → submit evaluation → fetch history.")
bullet(doc, "Test fallback: ASR_MODEL_DIR tidak diset → mock fallback aktif.")

heading(doc, "16.3 End-to-End Testing", level=2)
bullet(doc, "Setup docker-compose up (Postgres + backend) + flutter run di emulator.")
bullet(doc, "Skenario: register user baru → buka Jilid 1 Hal 1 → rekam 5 grup → cek progres.")
bullet(doc, "Skenario edge: rekam grup dengan audio kosong → expect incorrect/error gracefully.")
bullet(doc, "Skenario edge: putus koneksi saat evaluating → expect error message yang ramah.")

heading(doc, "16.4 User Acceptance Testing", level=2)
bullet(doc, "5-10 partisipan demo TA (anak SD + pemula dewasa).")
bullet(doc, "Tugas: ikuti onboarding, selesaikan 3 halaman Jilid 1, lihat progres.")
bullet(doc, "Metrik: System Usability Scale (SUS), task completion rate, time-to-complete.")
bullet(doc, "Feedback kualitatif: kemudahan navigasi, kejelasan visual feedback, kepercayaan terhadap evaluasi AI.")

heading(doc, "16.5 Audio Quality Testing", level=2)
bullet(doc, "Test rekam di lingkungan: tenang (< 30dB), normal (40-50dB), berisik (> 60dB).")
bullet(doc, "Verifikasi durasi rekaman 1-4 detik (sesuai data training).")
bullet(doc, "Test sample rate: pastikan output 16 kHz (cek dengan soundfile.read).")
bullet(doc, "Test orang berbeda: anak laki-laki, anak perempuan, dewasa pria, dewasa wanita.")

page_break(doc)

# ============ 17. STRATEGI DEPLOYMENT ============
heading(doc, "17. Strategi Deployment", level=1)

heading(doc, "17.1 Lingkungan Lokal Development", level=2)
code_block(doc, """\
# 1. Backend
cd Tilawati/backend
docker compose up -d   # PostgreSQL
uvicorn app.main:app --reload --port 8000

# 2. Frontend
cd Tilawati/frontend
flutter run            # emulator/device
""")
bullet(doc, "Backend default ASR_MODEL_DIR kosong → fallback mock aktif (cepat untuk dev UI).")
bullet(doc, "Set ASR_MODEL_DIR=Model/Model Final ASR/results (1)/asr_hpt_d untuk evaluasi real.")

heading(doc, "17.2 Staging", level=2)
bullet(doc, "Deploy backend ke server lab/komputer yang accessible LAN.")
bullet(doc, "Update ApiConfig.baseUrl ke http://<LAN-IP>:8000.")
bullet(doc, "Build APK debug → distribute ke partisipan UAT.")

heading(doc, "17.3 Production — HuggingFace Spaces (Backend)", level=2)
bullet(doc, "Buat Space tipe Docker.")
bullet(doc, "Upload Dockerfile + requirements + folder model HPT-D.")
bullet(doc, "Set ENV: ASR_MODEL_DIR, CLASSIFICATION_MODEL_DIR, DATABASE_URL, SECRET_KEY.")
bullet(doc, "Endpoint Spaces menjadi ASR_SERVICE_URL untuk failover dari instance lokal.")
bullet(doc, "Catatan: PostgreSQL terpisah (managed Postgres misal Supabase/Neon).")

heading(doc, "17.4 Distribusi Frontend", level=2)
bullet(doc, "Build APK release: flutter build apk --release.")
bullet(doc, "Build app bundle untuk Play Store: flutter build appbundle --release.")
bullet(doc, "Build iOS: flutter build ios --release (butuh macOS + Apple developer account).")
bullet(doc, "Build Web: flutter build web → host di Firebase Hosting / GitHub Pages.")

heading(doc, "17.5 Database Migration (Alembic)", level=2)
para(doc, "Saat ini schema dibuat via create_all(). Untuk produksi:")
code_block(doc, """\
cd Tilawati/backend
alembic init alembic
alembic revision --autogenerate -m "initial schema"
alembic upgrade head""")
bullet(doc, "Setiap perubahan model di app/models/ → buat revision baru.")
bullet(doc, "CI/CD trigger alembic upgrade head sebelum deploy aplikasi.")

heading(doc, "17.6 Monitoring & Logging", level=2)
bullet(doc, "FastAPI/uvicorn log default ke stdout — capture via Docker logs atau pipe ke ELK.")
bullet(doc, "Tambahkan request ID middleware untuk trace request lintas service.")
bullet(doc, "Health check endpoint /health untuk Liveness probe.")
bullet(doc, "Metrik: response time per endpoint, error rate, evaluasi sukses vs fallback.")

page_break(doc)

# ============ 18. RISIKO ============
heading(doc, "18. Risiko & Mitigasi", level=1)

make_table(doc,
    ["Kategori", "Risiko", "Dampak", "Mitigasi"],
    [
        ("Teknis-ML",
         "Akurasi HPT-D rendah untuk huruf TSA (65.2%)",
         "User frustrasi dengan false negative",
         "Gunakan fallback toleransi konfusi fonetik (TSA→SA dianggap partial); edukasi user di onboarding"),
        ("Teknis-ML",
         "Model HPT-D file 1.2 GB",
         "Ukuran APK besar jika di-bundle",
         "Letakkan model di backend (HuggingFace Spaces); APK tetap kecil"),
        ("Teknis-Backend",
         "Koneksi user 4G lambat",
         "Latensi evaluasi > 5 detik, UX buruk",
         "Tampilkan progress indicator; kompres audio sebelum upload jika perlu"),
        ("Teknis-Audio",
         "Mikrofon murah dengan SNR rendah",
         "Hasil ASR salah",
         "Tampilkan hint 'rekam di tempat sepi'; tampilkan waveform untuk feedback"),
        ("Data",
         "Audio user tersimpan di server",
         "Privacy concern",
         "Tambahkan policy retention; hapus audio submit_word segera; lesson submit retain hanya 30 hari"),
        ("Data",
         "Database backup tidak ada",
         "Kehilangan data evaluasi user",
         "Setup pg_dump scheduled backup; offsite copy mingguan"),
        ("Pengguna",
         "Anak frustrasi dengan kegagalan ASR",
         "Drop-out engagement",
         "Visual feedback positif kuat; izinkan retry tanpa penalty; mode 'tanpa skor'"),
        ("Pengguna",
         "Orangtua tidak paham hasil",
         "Tidak bisa mendukung anak",
         "Tambahkan halaman bantuan; tutorial video di pertama kali login"),
        ("Operasional",
         "HuggingFace Spaces down",
         "Aplikasi tidak bisa evaluasi",
         "Fallback mock 70%; cache konten lokal; UX message 'sedang offline'"),
        ("Akademik",
         "Validitas evaluasi tidak diakui pakar Tilawati",
         "Reputasi proyek/akademik",
         "Validasi dengan ustadz pakar; sertakan disclaimer 'AI sebagai alat bantu, bukan pengganti'"),
        ("Keamanan",
         "JWT SECRET_KEY bocor",
         "Token forge",
         "Rotate SECRET_KEY berkala; gunakan HSM/KMS di produksi"),
        ("Keamanan",
         "DDoS endpoint evaluasi (model intensif)",
         "Server overload",
         "Rate limiting per user (misal 100/menit); CDN frontfacing; circuit breaker"),
    ],
    col_widths=[1.0, 1.6, 1.4, 2.4],
)

page_break(doc)

# ============ 19. PENGEMBANGAN MASA DEPAN ============
heading(doc, "19. Pengembangan Masa Depan", level=1)

heading(doc, "19.1 Konten & Kurikulum", level=2)
bullet(doc, "Pengayaan konten Jilid 2-6 lengkap (200+ halaman total).")
bullet(doc, "Penambahan modul khusus do'a-do'a harian.")
bullet(doc, "Modul juz amma (juz 30) untuk pengguna lanjutan.")

heading(doc, "19.2 Mode Pembelajaran Baru", level=2)
bullet(doc, "Mode kuis glosarium: visual → audio (lihat huruf, ucapkan).")
bullet(doc, "Mode kuis glosarium: audio → visual (dengar audio referensi, pilih huruf).")
bullet(doc, "Mode tantangan: time-attack 28 huruf hijaiyah.")
bullet(doc, "Mode talqin: putar audio ustadz, user menirukan, AI evaluasi.")

heading(doc, "19.3 Personalisasi & AI", level=2)
bullet(doc, "Adaptive learning: identifikasi huruf yang sering salah, generate latihan ekstra otomatis.")
bullet(doc, "Recommendation engine: jika user di Jilid 2, rekomendasikan halaman Jilid 1 yang harus diulang.")
bullet(doc, "Fine-tune model HPT-D dengan data user (federated learning) untuk personalisasi pelafalan.")

heading(doc, "19.4 Kolaborasi & Komunitas", level=2)
bullet(doc, "Akun ustadz/lembaga: bisa monitor murid, beri komentar, atur kurikulum.")
bullet(doc, "Group challenge mingguan untuk TPQ.")
bullet(doc, "Leaderboard friendly (anonim, opsional).")

heading(doc, "19.5 Aksesibilitas & Inklusivitas", level=2)
bullet(doc, "Multi-bahasa: Bahasa Indonesia, Inggris, Arab.")
bullet(doc, "Mode tanpa suara: visual-only untuk pengguna tuli (membaca dengan jari di mulut, dll).")
bullet(doc, "Font size adjustable untuk lansia.")
bullet(doc, "Tema gelap.")

heading(doc, "19.6 Konten & Media Tambahan", level=2)
bullet(doc, "Audio referensi tiap huruf (rekaman ustadz).")
bullet(doc, "Video tutorial makharijul huruf (lokasi titik artikulasi di mulut).")
bullet(doc, "Animasi gerak bibir untuk huruf yang mirip (ث vs س, ح vs ه).")

heading(doc, "19.7 Backend & Infrastruktur", level=2)
bullet(doc, "Migrasi ke microservices: auth-service, evaluation-service, content-service.")
bullet(doc, "Kafka/RabbitMQ untuk async submit (high-volume).")
bullet(doc, "GraphQL API alternatif untuk client lebih fleksibel.")
bullet(doc, "Edge inference: TFLite/ONNX export model HPT-D untuk inference on-device.")

page_break(doc)

# ============ 20. LAMPIRAN ============
heading(doc, "20. Lampiran", level=1)

heading(doc, "Lampiran A. Daftar Konten Tilawati Jilid 1 (44 Halaman)", level=2)
make_table(doc,
    ["Hal", "Topik"],
    [
        ("1", "Huruf Alif & Ba"),
        ("2", "Huruf Ta"),
        ("3", "Huruf Tsa"),
        ("4", "Huruf Jim & Ha"),
        ("5", "Huruf Kha"),
        ("6", "Latihan kombinasi"),
        ("7", "Huruf Dal"),
        ("8", "Huruf Dzal"),
        ("9", "Huruf Ra & Za"),
        ("10", "Huruf Sin"),
        ("11", "Huruf Syin"),
        ("12", "Latihan kombinasi"),
        ("13", "Huruf Shad"),
        ("14", "Huruf Dhad"),
        ("15", "Huruf Tha"),
        ("16", "Huruf Zha"),
        ("17", "Huruf 'Ain"),
        ("18", "Latihan kombinasi"),
        ("19", "Huruf Ghain"),
        ("20", "Huruf Fa"),
        ("21", "Huruf Qaf"),
        ("22", "Latihan kombinasi"),
        ("23", "Huruf Kaf"),
        ("24", "Huruf Lam"),
        ("25", "Huruf Mim"),
        ("26", "Huruf Nun"),
        ("27", "Huruf Wawu"),
        ("28", "Huruf Ha"),
        ("29", "Latihan kombinasi"),
        ("30", "Huruf Hamzah"),
        ("31", "Huruf Ya"),
        ("32", "Latihan 1"),
        ("33", "Latihan 9 (BA TA TSA)"),
        ("34", "Latihan 10 (JA HA KHA)"),
        ("35", "Latihan 11 (SA SYA)"),
        ("36", "Latihan 12 (SHA DHA)"),
        ("37", "Latihan 13 ('A GHA)"),
        ("38", "Latihan 14 (FA QA)"),
        ("39", "Latihan 15 (LA)"),
        ("40", "Latihan 16 (KA)"),
        ("41", "Latihan 17 (NA)"),
        ("42", "Latihan 18 (HA)"),
        ("43", "Latihan 19 (Hamzah-Ya)"),
        ("44", "Latihan 20 (Kombinasi penuh)"),
    ],
    col_widths=[0.6, 5.0],
)

heading(doc, "Lampiran B. Skema Penilaian Detail", level=2)
make_table(doc,
    ["Komponen", "Bobot", "Cara Hitung"],
    [
        ("Makharijul Huruf", "40%", "Persentase token ASR yang match dengan expected"),
        ("Tajwid", "35%", "Placeholder saat ini; ke depan integrasi model tajwid"),
        ("Kelancaran", "25%", "Berdasarkan ratio durasi audio terhadap jumlah huruf yang diharapkan"),
    ],
    col_widths=[1.6, 0.8, 4.0],
)

para(doc, "Formula:")
code_block(doc, """\
overall_score = 0.40 * makharijul + 0.35 * tajwid + 0.25 * kelancaran

Grade:
  ≥ 85   → Mumtaz       (Sangat Baik)
  70-84  → Jayyid Jiddan (Baik Sekali)
  55-69  → Jayyid       (Baik)
  < 55   → Maqbul       (Cukup)
""")

heading(doc, "Lampiran C. Glossary Istilah", level=2)
make_table(doc,
    ["Istilah", "Definisi"],
    [
        ("Tilawati", "Metode pembelajaran membaca Al-Qur'an asal Surabaya, terdiri dari 6 jilid."),
        ("Hijaiyah", "28 huruf alfabet bahasa Arab yang menjadi unit dasar Al-Qur'an."),
        ("Fathah", "Harakat (tanda baca) yang melambangkan vokal 'a' di atas huruf."),
        ("Makharijul Huruf", "Titik artikulasi huruf — tempat keluarnya huruf dari mulut/tenggorokan."),
        ("Tajwid", "Aturan pelafalan Al-Qur'an (panjang pendek, dengung, samar, dll)."),
        ("Mumtaz", "Predikat 'sangat baik' dalam evaluasi (sering ≥ 85%)."),
        ("Jayyid Jiddan", "Predikat 'baik sekali' (70-84%)."),
        ("Jayyid", "Predikat 'baik' (55-69%)."),
        ("Maqbul", "Predikat 'cukup' (<55%)."),
        ("Talqin", "Mencontohkan bacaan oleh ustadz untuk ditirukan murid."),
        ("ASR", "Automatic Speech Recognition — sistem AI pengubah ucapan ke teks."),
        ("CTC", "Connectionist Temporal Classification — loss function untuk ASR variable-length."),
        ("CER/WER", "Character/Word Error Rate — metrik akurasi ASR (semakin rendah semakin baik)."),
        ("TSA", "Token-level Sequence Accuracy — metrik akurasi rangkaian token."),
        ("HPT-D", "Hyperparameter Tuning variant D — model ASR final TilawatiMobile."),
        ("Wav2Vec2", "Arsitektur self-supervised audio dari Meta AI, basis model TilawatiMobile."),
        ("XLS-R", "Cross-Lingual Speech Representation — varian Wav2Vec2 multibahasa."),
        ("HPT", "HyperParameter Tuning — proses eksperimen sistematis untuk mencari konfigurasi optimal."),
    ],
    col_widths=[1.6, 4.8],
)

heading(doc, "Lampiran D. Referensi", level=2)
bullet(doc, "Tilawati: Metode Praktis Belajar Al-Qur'an. Tim Tilawati Pesantren Virtual An-Nur.")
bullet(doc, "Conneau, A., et al. (2020). Wav2Vec2-XLS-R: Unsupervised Cross-Lingual Representation Learning for Speech Recognition. Meta AI.")
bullet(doc, "anas/wav2vec2-large-xlsr-arabic — HuggingFace Model Card.")
bullet(doc, "Graves, A. (2006). Connectionist Temporal Classification: Labelling Unsegmented Sequence Data with Recurrent Neural Networks. ICML 2006.")
bullet(doc, "FastAPI Documentation — fastapi.tiangolo.com")
bullet(doc, "Flutter Documentation — docs.flutter.dev")
bullet(doc, "Dokumen TA: Blueprint_ModelASR_tilawati.docx (Spesifikasi teknis Phase 2 ASR)")
bullet(doc, "Dokumen TA: rancangan_deployment.md (Rancangan arsitektur deployment Phase 3)")
bullet(doc, "Dokumen TA: asr_hpt_d_analysis.md (Post-mortem model final HPT-D)")

# ============ FOOTER PAGE ============
page_break(doc)
heading(doc, "Penutup", level=1)
para(doc,
    "Dokumen blueprint ini menyajikan gambaran komprehensif rancangan pengembangan TilawatiMobile "
    "secara end-to-end — dari komponen AI/ML, backend, frontend, hingga roadmap dan strategi deployment. "
    "Dokumen ini bersifat hidup (living document) dan akan diperbarui seiring perkembangan proyek. "
    "Semua keputusan arsitektural, status implementasi, dan rencana ke depan berbasis pada kondisi "
    "kode aktual per 30 April 2026."
)

para(doc,
    "Tujuan akhir TilawatiMobile bukan hanya menjadi proyek akademik, tetapi diharapkan dapat "
    "memberikan kontribusi nyata pada ekosistem pembelajaran Al-Qur'an di Indonesia — sebagai alat "
    "bantu (bukan pengganti) ustadz/ustadzah untuk mendemokratisasi akses koreksi bacaan yang akurat "
    "dan terjangkau."
)

doc.add_paragraph()
sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sig.add_run("— Selesai —")
r.italic = True
r.font.color.rgb = GRAY
r.font.size = Pt(10)

# ============ SAVE ============
output_path = 'blueprint_TilawatiMobile.docx'
doc.save(output_path)
print(f"[OK] Blueprint berhasil dibuat: {output_path}")
